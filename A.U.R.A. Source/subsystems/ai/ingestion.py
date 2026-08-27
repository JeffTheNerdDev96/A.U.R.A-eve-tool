# -*- coding: utf-8 -*-
# ==============================================================================
# Adaptive Underworld Recon Array (A.U.R.A.)
# Copyright (C) 2026 JeffTheNerdDev96
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.
# ==============================================================================
"""
Multiformat Tactical Ingestion Engine for EVE Online Adaptive Underworld Recon Array (A.U.R.A.).
Handles in-game screenshots (Overviews, D-Scans, Killmails), EFT fits, logs, and tactical briefs.
"""
import os
from typing import Dict, Any, List
from PIL import Image, ImageEnhance

from core.config import config
from core.error_handler import AURAErrorCode, log_diagnostic_error
from core.input_safety import clamp_text, strip_control_chars


def _check_file_size(file_path: str) -> None:
    size = os.path.getsize(file_path)
    if size > config.max_attachment_bytes:
        raise ValueError(
            f"File exceeds {config.max_attachment_bytes // (1024 * 1024)} MB attachment limit"
        )


def _finalize_text(text: str) -> str:
    return clamp_text(strip_control_chars(text or ""), config.max_llm_context_chars // 2)


class ImagePreprocessor:
    """Hardware-accelerated visual reader and neural OCR processor for EVE screenshots."""
    
    @staticmethod
    def analyze_image_content(image_path: str) -> Dict[str, Any]:
        img = None
        scaled_img = None
        enhanced_img = None
        try:
            _check_file_size(image_path)
            Image.MAX_IMAGE_PIXELS = config.max_image_pixels
            img = Image.open(image_path).convert("RGB")
            w, h = img.size
            extracted_lines = []
            
            scale_factor = 1.5
            if w > 1920 or h > 1080:
                scale_factor = min(1920 / w, 1080 / h)
            
            new_w = max(320, int(w * scale_factor))
            new_h = max(240, int(h * scale_factor))
            
            scaled_img = img.resize((new_w, new_h), Image.Resampling.BILINEAR)
            enhancer = ImageEnhance.Contrast(scaled_img)
            enhanced_img = enhancer.enhance(1.3)
            
            try:
                import winocr
                res = winocr.recognize_pil_sync(enhanced_img, 'en')
                if res and "lines" in res:
                    for line_dict in res["lines"]:
                        txt = line_dict.get("text", "").strip()
                        if txt and txt not in extracted_lines:
                            extracted_lines.append(txt)
                elif res and "text" in res:
                    for line in res["text"].split("\n"):
                        clean = line.strip()
                        if clean and clean not in extracted_lines:
                            extracted_lines.append(clean)
            except Exception:
                pass

            extracted_text = _finalize_text(
                "\n".join(extracted_lines) if extracted_lines
                else "[Screenshot contains visual tactical graphics without distinct machine-readable text]"
            )
            
            return {
                "dimensions": f"{w}x{h}",
                "extracted_text": extracted_text,
                "text_lines": extracted_lines,
                "has_text": len(extracted_lines) > 0,
                "summary": f"Tactical Screenshot ({w}x{h}) - {len(extracted_lines)} telemetry element(s) detected"
            }
        except Exception as e:
            return {
                "dimensions": "Unknown",
                "extracted_text": f"[Error analyzing tactical image: {e}]",
                "text_lines": [],
                "has_text": False,
                "summary": f"Error: {e}"
            }
        finally:
            if enhanced_img:
                try: enhanced_img.close()
                except Exception: pass
            if scaled_img:
                try: scaled_img.close()
                except Exception: pass
            if img:
                try: img.close()
                except Exception: pass



class DocumentParser:
    """Universal parser for EVE Online fleet briefs, chat logs, fits, and documents."""
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        try:
            import pypdf
        except ImportError as e:
            return f"[Error: pypdf library not installed: {e}]"
        try:
            reader = pypdf.PdfReader(file_path)
            pages_text = []
            for idx, page in enumerate(reader.pages):
                if idx >= config.max_pdf_pages:
                    pages_text.append(f"[Truncated at {config.max_pdf_pages} pages]")
                    break
                text = page.extract_text()
                if text and text.strip():
                    pages_text.append(f"--- Page {idx + 1} ---\n{text.strip()}")
            return _finalize_text("\n\n".join(pages_text) if pages_text else "[PDF appears empty]")
        except Exception as e:
            return f"[Error parsing PDF: {e}]"

    @staticmethod
    def parse_docx(file_path: str) -> str:
        try:
            import docx
        except ImportError as e:
            return f"[Error: python-docx library not installed: {e}]"
        try:
            doc = docx.Document(file_path)
            content = []
            para_count = 0
            for p in doc.paragraphs:
                if para_count >= config.max_docx_paragraphs:
                    content.append("[Truncated: paragraph limit reached]")
                    break
                if p.text.strip():
                    content.append(p.text.strip())
                    para_count += 1
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        table_data.append(" | ".join(row_text))
                if table_data:
                    content.append("\n[Table]\n" + "\n".join(table_data))
            return _finalize_text("\n\n".join(content) if content else "[Word Document is empty]")
        except Exception as e:
            return f"[Error parsing DOCX: {e}]"

    @staticmethod
    def _read_text_file_safe(file_path: str) -> str:
        """Reads tactical text files with robust multi-encoding fallbacks."""
        max_bytes = config.max_attachment_bytes
        try:
            with open(file_path, "rb") as raw:
                data = raw.read(max_bytes + 1)
            if len(data) > max_bytes:
                raise ValueError(f"File exceeds {max_bytes // (1024 * 1024)} MB attachment limit")
        except OSError as e:
            return f"[Error reading file: {e}]"
        for enc in ["utf-8-sig", "utf-8", "cp1252", "latin-1"]:
            try:
                return _finalize_text(data.decode(enc))
            except (UnicodeDecodeError, LookupError):
                continue
        return _finalize_text(data.decode("utf-8", errors="ignore"))

    @staticmethod
    def parse_file(file_path: str) -> Dict[str, Any]:
        try:
            return DocumentParser._parse_file_impl(file_path)
        except Exception as exc:
            log_diagnostic_error(
                AURAErrorCode.ERR_3004_INGESTION_FAILED,
                exc,
                f"DocumentParser.parse_file({file_path})",
            )
            filename = os.path.basename(file_path)
            return {
                "filename": filename,
                "path": file_path,
                "type": "error",
                "text": f"[Ingestion failed: {exc}]",
                "summary": f"Ingestion failed for {filename}",
            }

    @staticmethod
    def _parse_file_impl(file_path: str) -> Dict[str, Any]:
        _check_file_size(file_path)
        filename = os.path.basename(file_path)
        ext = os.path.splitext(file_path)[1].lower()
        
        match ext:
            case ".png" | ".jpg" | ".jpeg" | ".bmp" | ".webp":
                analysis = ImagePreprocessor.analyze_image_content(file_path)
                return {
                    "filename": filename,
                    "path": file_path,
                    "type": "image",
                    "text": analysis["extracted_text"],
                    "analysis": analysis,
                    "summary": analysis["summary"]
                }
                
            case ".pdf":
                text = DocumentParser.parse_pdf(file_path)
                word_est = max(1, len(text) // 6)
                return {
                    "filename": filename,
                    "path": file_path,
                    "type": "document",
                    "text": text,
                    "summary": f"PDF fleet doctrine (~{word_est:,} words)"
                }
                
            case ".docx" | ".doc":
                text = DocumentParser.parse_docx(file_path)
                word_est = max(1, len(text) // 6)
                return {
                    "filename": filename,
                    "path": file_path,
                    "type": "document",
                    "text": text,
                    "summary": f"Word document (~{word_est:,} words)"
                }
                
            case ".txt" | ".csv" | ".md" | ".json" | ".log" | ".eft" | ".xml":
                text = DocumentParser._read_text_file_safe(file_path)
                word_est = max(1, len(text) // 6)
                return {
                    "filename": filename,
                    "path": file_path,
                    "type": "document",
                    "text": text,
                    "summary": f"Tactical text file (~{word_est:,} words)"
                }
                
            case _:
                return {
                    "filename": filename,
                    "path": file_path,
                    "type": "unknown",
                    "text": f"[Unsupported format: {ext}]",
                    "summary": f"Unsupported format: {ext}"
                }
