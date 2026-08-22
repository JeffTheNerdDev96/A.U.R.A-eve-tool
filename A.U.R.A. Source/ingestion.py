"""
Multiformat Tactical Ingestion Engine for EVE Online Adaptive Underworld Recon Array (A.U.R.A.).
Handles in-game screenshots (Overviews, D-Scans, Killmails), EFT fits, logs, and tactical briefs.
"""
import os
from typing import Dict, Any, List
from PIL import Image, ImageEnhance

try:
    import llama_cpp
except Exception:
    pass


class ImagePreprocessor:
    """Hardware-accelerated visual reader and neural OCR processor for EVE screenshots."""
    
    @staticmethod
    def analyze_image_content(image_path: str) -> Dict[str, Any]:
        img = None
        scaled_img = None
        enhanced_img = None
        try:
            img = Image.open(image_path).convert("RGB")
            w, h = img.size
            extracted_lines = []
            
            # Bound dimensions to max 1920 to prevent excessive RAM allocation on 4K screenshots
            scale_factor = 1.5
            if w > 1920 or h > 1080:
                scale_factor = min(1920 / w, 1080 / h)
            
            new_w = max(320, int(w * scale_factor))
            new_h = max(240, int(h * scale_factor))
            
            scaled_img = img.resize((new_w, new_h), Image.Resampling.BILINEAR)
            enhancer = ImageEnhance.Contrast(scaled_img)
            enhanced_img = enhancer.enhance(1.3)
            
            try:
                import winocr
                res = winocr.recognize_pil_sync(enhanced_img, 'en')
                if res and "lines" in res:
                    for line_dict in res["lines"]:
                        txt = line_dict.get("text", "").strip()
                        if txt and txt not in extracted_lines:
                            extracted_lines.append(txt)
                elif res and "text" in res:
                    for line in res["text"].split("\n"):
                        clean = line.strip()
                        if clean and clean not in extracted_lines:
                            extracted_lines.append(clean)
            except Exception:
                pass

            extracted_text = "\n".join(extracted_lines) if extracted_lines else "[Screenshot contains visual tactical graphics without distinct machine-readable text]"
            
            return {
                "dimensions": f"{w}x{h}",
                "extracted_text": extracted_text,
                "text_lines": extracted_lines,
                "has_text": len(extracted_lines) > 0,
                "summary": f"Tactical Screenshot ({w}x{h}) - {len(extracted_lines)} telemetry element(s) detected"
            }
        except Exception as e:
            return {
                "dimensions": "Unknown",
                "extracted_text": f"[Error analyzing tactical image: {e}]",
                "text_lines": [],
                "has_text": False,
                "summary": f"Error: {e}"
            }
        finally:
            if enhanced_img:
                try: enhanced_img.close()
                except Exception: pass
            if scaled_img:
                try: scaled_img.close()
                except Exception: pass
            if img:
                try: img.close()
                except Exception: pass



class DocumentParser:
    """Universal parser for EVE Online fleet briefs, chat logs, fits, and documents."""
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        try:
            import pypdf
        except ImportError as e:
            return f"[Error: pypdf library not installed: {e}]"
        try:
            reader = pypdf.PdfReader(file_path)
            pages_text = []
            for idx, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages_text.append(f"--- Page {idx + 1} ---\n{text.strip()}")
            return "\n\n".join(pages_text) if pages_text else "[PDF appears empty]"
        except Exception as e:
            return f"[Error parsing PDF: {e}]"

    @staticmethod
    def parse_docx(file_path: str) -> str:
        try:
            import docx
        except ImportError as e:
            return f"[Error: python-docx library not installed: {e}]"
        try:
            doc = docx.Document(file_path)
            content = []
            for p in doc.paragraphs:
                if p.text.strip():
                    content.append(p.text.strip())
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        table_data.append(" | ".join(row_text))
                if table_data:
                    content.append("\n[Table]\n" + "\n".join(table_data))
            return "\n\n".join(content) if content else "[Word Document is empty]"
        except Exception as e:
            return f"[Error parsing DOCX: {e}]"

    @staticmethod
    def _read_text_file_safe(file_path: str) -> str:
        """Reads tactical text files with robust multi-encoding fallbacks."""
        for enc in ["utf-8-sig", "utf-8", "cp1252", "latin-1"]:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, LookupError):
                continue
            except Exception as e:
                return f"[Error reading file: {e}]"
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            return f"[Error reading file: {e}]"

    @staticmethod
    def parse_file(file_path: str) -> Dict[str, Any]:
        filename = os.path.basename(file_path)
        ext = os.path.splitext(file_path)[1].lower()
        
        match ext:
            # 1. Images / Screenshots
            case ".png" | ".jpg" | ".jpeg" | ".bmp" | ".webp":
                analysis = ImagePreprocessor.analyze_image_content(file_path)
                return {
                    "filename": filename,
                    "path": file_path,
                    "type": "image",
                    "text": analysis["extracted_text"],
                    "analysis": analysis,
                    "summary": analysis["summary"]
                }
                
            # 2. PDF
            case ".pdf":
                text = DocumentParser.parse_pdf(file_path)
                word_est = max(1, len(text) // 6)
                return {
                    "filename": filename,
                    "path": file_path,
                    "type": "document",
                    "text": text,
                    "summary": f"PDF fleet doctrine (~{word_est:,} words)"
                }
                
            # 3. Word DOCX
            case ".docx" | ".doc":
                text = DocumentParser.parse_docx(file_path)
                word_est = max(1, len(text) // 6)
                return {
                    "filename": filename,
                    "path": file_path,
                    "type": "document",
                    "text": text,
                    "summary": f"Word document (~{word_est:,} words)"
                }
                
            # 4. Text / EFT Fits / Logs / CSV
            case ".txt" | ".csv" | ".md" | ".json" | ".log" | ".py" | ".eft" | ".xml":
                text = DocumentParser._read_text_file_safe(file_path)
                word_est = max(1, len(text) // 6)
                return {
                    "filename": filename,
                    "path": file_path,
                    "type": "document",
                    "text": text,
                    "summary": f"Tactical text file (~{word_est:,} words)"
                }
                
            case _:
                return {
                    "filename": filename,
                    "path": file_path,
                    "type": "unknown",
                    "text": f"[Unsupported format: {ext}]",
                    "summary": f"Unsupported format: {ext}"
                }
